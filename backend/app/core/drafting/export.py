"""Export legal documents to DOCX and PDF formats.

Provides async export functions that produce properly formatted Indian legal
documents with Times New Roman typography, 1-inch margins, and structured
section headings.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone
from typing import TYPE_CHECKING

from docx import Document  # type: ignore[import-untyped]
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[import-untyped]
from docx.shared import Inches, Pt  # type: ignore[import-untyped]
from reportlab.lib.pagesizes import A4, legal  # type: ignore[import-untyped]
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet  # type: ignore[import-untyped]
from reportlab.lib.units import cm, inch  # type: ignore[import-untyped]
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer  # type: ignore[import-untyped]

from app.core.drafting.templates import DocumentTemplate

if TYPE_CHECKING:
    from app.core.drafting.court_profiles import CourtProfile

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _is_heading(line: str) -> bool:
    """Return True if *line* looks like a section heading."""
    stripped = line.strip()
    if stripped.startswith("#"):
        return True
    # All-caps lines of 4+ characters (e.g. "PRAYER", "FACTS OF THE CASE")
    if len(stripped) >= 4 and stripped == stripped.upper() and stripped.replace(" ", "").isalpha():
        return True
    return False


def _clean_heading(line: str) -> str:
    """Strip markdown heading markers and return clean text."""
    return line.lstrip("#").strip()


def _parse_sections(content: str) -> list[tuple[str | None, list[str]]]:
    """Parse *content* into (heading, body_lines) pairs.

    A ``None`` heading means body text before the first heading.
    """
    sections: list[tuple[str | None, list[str]]] = []
    current_heading: str | None = None
    current_body: list[str] = []

    for line in content.split("\n"):
        if _is_heading(line):
            # Flush previous section
            sections.append((current_heading, current_body))
            current_heading = _clean_heading(line)
            current_body = []
        else:
            current_body.append(line)

    # Flush last section
    sections.append((current_heading, current_body))
    return sections


# ---------------------------------------------------------------------------
# DOCX export
# ---------------------------------------------------------------------------


async def export_to_docx(
    content: str,
    template: DocumentTemplate,
    *,
    title: str = "",
    court_profile: CourtProfile | None = None,
    affidavit: str = "",
) -> bytes:
    """Export document content to DOCX format with Indian legal formatting.

    Args:
        content: The full document text (may contain markdown headings).
        template: The DocumentTemplate describing the document type.
        title: Optional title override; defaults to ``template.display_name``.
        court_profile: Optional court-specific formatting profile. When
            provided, margins, font sizes and line spacing are taken from the
            profile instead of the defaults.
        affidavit: Optional affidavit content to append after the main
            document (separated by a page break).

    Returns:
        The DOCX file as raw bytes.
    """
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        if court_profile is not None:
            section.top_margin = Inches(court_profile.margin_top_cm / 2.54)
            section.bottom_margin = Inches(court_profile.margin_bottom_cm / 2.54)
            section.left_margin = Inches(court_profile.margin_left_cm / 2.54)
            section.right_margin = Inches(court_profile.margin_right_cm / 2.54)
        else:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    # -- Resolve font sizes from profile or defaults --
    body_font_size = Pt(court_profile.font_size_body) if court_profile else Pt(12)
    heading_font_size = Pt(court_profile.font_size_heading) if court_profile else Pt(14)
    line_spacing = court_profile.line_spacing if court_profile else None

    # -- Default font --
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = body_font_size

    # -- Document title --
    doc_title = title or template.display_name
    title_para = doc.add_heading(doc_title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True

    # -- Parse and render sections --
    paragraph_number = 1
    sections = _parse_sections(content)

    for heading, body_lines in sections:
        if heading is not None:
            heading_para = doc.add_heading(heading, level=1)
            for run in heading_para.runs:
                run.font.name = "Times New Roman"
                run.font.size = heading_font_size
                run.font.bold = True

        body_text = "\n".join(body_lines).strip()
        if not body_text:
            continue

        for paragraph_text in body_text.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue

            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            if line_spacing is not None:
                para.paragraph_format.line_spacing = line_spacing

            # Add paragraph number for substantive content
            if heading is not None:
                run = para.add_run(f"{paragraph_number}. ")
                run.font.name = "Times New Roman"
                run.font.size = body_font_size
                run.font.bold = True
                paragraph_number += 1

            run = para.add_run(paragraph_text.replace("\n", " "))
            run.font.name = "Times New Roman"
            run.font.size = body_font_size

    # -- Affidavit section --
    if affidavit:
        doc.add_page_break()
        aff_title = doc.add_heading("AFFIDAVIT", level=0)
        aff_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in aff_title.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(16)
            run.font.bold = True

        aff_sections = _parse_sections(affidavit)
        for heading, body_lines in aff_sections:
            if heading is not None:
                heading_para = doc.add_heading(heading, level=1)
                for run in heading_para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = heading_font_size
                    run.font.bold = True

            body_text = "\n".join(body_lines).strip()
            if not body_text:
                continue

            for paragraph_text in body_text.split("\n\n"):
                paragraph_text = paragraph_text.strip()
                if not paragraph_text:
                    continue

                para = doc.add_paragraph()
                para.paragraph_format.space_after = Pt(6)
                if line_spacing is not None:
                    para.paragraph_format.line_spacing = line_spacing

                run = para.add_run(paragraph_text.replace("\n", " "))
                run.font.name = "Times New Roman"
                run.font.size = body_font_size

    # -- Document metadata --
    doc.core_properties.title = doc_title
    doc.core_properties.author = "Smriti AI"
    doc.core_properties.created = datetime.now(timezone.utc)
    doc.core_properties.comments = (
        f"Generated by Smriti AI — {template.display_name}"
    )

    # -- Serialise to bytes --
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------


async def export_to_pdf(
    content: str,
    template: DocumentTemplate,
    *,
    title: str = "",
    court_profile: CourtProfile | None = None,
    affidavit: str = "",
) -> bytes:
    """Export document content to PDF format with Indian legal formatting.

    Args:
        content: The full document text (may contain markdown headings).
        template: The DocumentTemplate describing the document type.
        title: Optional title override; defaults to ``template.display_name``.
        court_profile: Optional court-specific formatting profile. When
            provided, margins, font sizes, line spacing and paper size are
            taken from the profile instead of the defaults.
        affidavit: Optional affidavit content to append after the main
            document (separated by a page break).

    Returns:
        The PDF file as raw bytes.
    """
    buf = io.BytesIO()

    doc_title = title or template.display_name

    # -- Resolve layout from profile or defaults --
    if court_profile is not None:
        pagesize = legal if court_profile.paper_size == "legal" else A4
        top_margin = court_profile.margin_top_cm * cm
        bottom_margin = court_profile.margin_bottom_cm * cm
        left_margin = court_profile.margin_left_cm * cm
        right_margin = court_profile.margin_right_cm * cm
        body_font_size = court_profile.font_size_body
        heading_font_size = court_profile.font_size_heading
        body_leading = court_profile.font_size_body * court_profile.line_spacing
    else:
        pagesize = A4
        top_margin = 1 * inch
        bottom_margin = 1 * inch
        left_margin = 1 * inch
        right_margin = 1 * inch
        body_font_size = 12
        heading_font_size = 14
        body_leading = 16

    doc = SimpleDocTemplate(
        buf,
        pagesize=pagesize,
        topMargin=top_margin,
        bottomMargin=bottom_margin,
        leftMargin=left_margin,
        rightMargin=right_margin,
        title=doc_title,
        author="Smriti AI",
        creator="Smriti AI",
        subject=f"Legal Document — {template.display_name}",
    )

    # -- Styles --
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=16,
        alignment=1,  # centre
        spaceAfter=20,
    )

    heading_style = ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=heading_font_size,
        spaceBefore=14,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        "BodyText",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=body_font_size,
        leading=body_leading,
        spaceAfter=6,
    )

    # -- Build flowable list --
    flowables: list[Paragraph | Spacer | PageBreak] = []

    flowables.append(Paragraph(doc_title, title_style))
    flowables.append(Spacer(1, 12))

    paragraph_number = 1
    sections = _parse_sections(content)

    for heading, body_lines in sections:
        if heading is not None:
            flowables.append(Paragraph(heading, heading_style))

        body_text = "\n".join(body_lines).strip()
        if not body_text:
            continue

        for paragraph_text in body_text.split("\n\n"):
            paragraph_text = paragraph_text.strip()
            if not paragraph_text:
                continue

            # Escape XML-sensitive characters for ReportLab
            safe_text = (
                paragraph_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )

            if heading is not None:
                safe_text = f"<b>{paragraph_number}.</b> {safe_text}"
                paragraph_number += 1

            flowables.append(Paragraph(safe_text, body_style))

    # -- Affidavit section --
    if affidavit:
        flowables.append(PageBreak())

        aff_title_style = ParagraphStyle(
            "AffTitle",
            parent=styles["Title"],
            fontName="Times-Bold",
            fontSize=16,
            alignment=1,
            spaceAfter=20,
        )
        flowables.append(Paragraph("AFFIDAVIT", aff_title_style))
        flowables.append(Spacer(1, 12))

        aff_sections = _parse_sections(affidavit)
        for heading, body_lines in aff_sections:
            if heading is not None:
                flowables.append(Paragraph(heading, heading_style))

            body_text = "\n".join(body_lines).strip()
            if not body_text:
                continue

            for paragraph_text in body_text.split("\n\n"):
                paragraph_text = paragraph_text.strip()
                if not paragraph_text:
                    continue

                safe_text = (
                    paragraph_text.replace("&", "&amp;")
                    .replace("<", "&lt;")
                    .replace(">", "&gt;")
                    .replace("\n", "<br/>")
                )
                flowables.append(Paragraph(safe_text, body_style))

    doc.build(flowables)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Research memo export (no template needed)
# ---------------------------------------------------------------------------


_INLINE_RE = __import__("re").compile(r"(\*\*[^\*]+\*\*|\*[^\*]+\*|`[^`]+`)")


def _add_inline_runs(paragraph, text: str, base_size: int = 12) -> None:
    """Add runs to a docx paragraph, parsing **bold**, *italic*, and `code`."""
    parts = _INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(base_size)
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run.text = part[2:-2]
            run.font.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run.text = part[1:-1]
            run.font.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run.text = part[1:-1]
            run.font.name = "Courier New"
        else:
            run.text = part


def _render_markdown_to_docx(doc, content: str, body_size: int = 12, heading_size: int = 14) -> None:
    """Render markdown content into a docx Document.

    Handles: # headings, **bold**, *italic*, | tables |, - bullets, 1. numbered,
    paragraph breaks, and <br> line breaks.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Heading
        if stripped.startswith("#"):
            level = 0
            while level < len(stripped) and stripped[level] == "#":
                level += 1
            text = stripped[level:].strip()
            level = min(max(level, 1), 4)
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(heading_size if level <= 2 else body_size + 1)
                run.font.bold = True
            i += 1
            continue

        # Horizontal rule
        if stripped in ("---", "***", "___"):
            doc.add_paragraph("_" * 60)
            i += 1
            continue

        # Table (starts with |, next line is separator |---|---|)
        if stripped.startswith("|") and i + 1 < len(lines) and "---" in lines[i + 1]:
            header_cells = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # Skip header and separator
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row_cells)
                i += 1
            if header_cells:
                table = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
                table.style = "Light Grid Accent 1"
                for idx, hc in enumerate(header_cells):
                    cell = table.rows[0].cells[idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(hc.replace("<br>", " "))
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(body_size - 1)
                    run.font.bold = True
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < len(header_cells):
                            cell = table.rows[r_idx + 1].cells[c_idx]
                            cell.text = ""
                            for sub_line in cell_text.split("<br>"):
                                p = cell.paragraphs[0] if sub_line == cell_text.split("<br>")[0] else cell.add_paragraph()
                                _add_inline_runs(p, sub_line.strip(), base_size=body_size - 1)
                doc.add_paragraph()  # Spacer
            continue

        # Bullet list
        if stripped.startswith(("- ", "* ", "+ ")):
            while i < len(lines) and lines[i].strip().startswith(("- ", "* ", "+ ")):
                item_text = lines[i].strip()[2:]
                para = doc.add_paragraph(style="List Bullet")
                _add_inline_runs(para, item_text, base_size=body_size)
                i += 1
            continue

        # Numbered list
        if stripped[0].isdigit() and ". " in stripped[:5]:
            while i < len(lines) and lines[i].strip() and lines[i].strip()[0:1].isdigit() and ". " in lines[i].strip()[:5]:
                item_text = lines[i].strip().split(". ", 1)[1] if ". " in lines[i].strip() else lines[i].strip()
                para = doc.add_paragraph(style="List Number")
                _add_inline_runs(para, item_text, base_size=body_size)
                i += 1
            continue

        # Regular paragraph — collect until blank line
        para_lines = []
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(("#", "|", "- ", "* ", "+ ")):
            s = lines[i].strip()
            if s[0:1].isdigit() and ". " in s[:5]:
                break
            para_lines.append(s)
            i += 1
        if para_lines:
            para_text = " ".join(para_lines).replace("<br>", " ")
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(6)
            _add_inline_runs(para, para_text, base_size=body_size)


async def export_research_memo_docx(
    content: str,
    *,
    title: str = "Research Memo",
    footnotes: list[dict] | None = None,
) -> bytes:
    """Export a research memo to DOCX with footnotes as bibliography."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(16)
        run.font.bold = True

    # Body with proper markdown rendering
    _render_markdown_to_docx(doc, content, body_size=12, heading_size=14)

    # Bibliography from footnotes
    if footnotes:
        doc.add_heading("Bibliography", level=1)
        for fn in footnotes:
            num = fn.get("number", 0)
            citation = fn.get("citation", "")
            title_text = fn.get("title", "")
            label = f"[{num}] {citation}"
            if title_text and title_text != citation:
                label += f" — {title_text}"
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(label)
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    doc.core_properties.title = title
    doc.core_properties.author = "Smriti AI"
    doc.core_properties.created = datetime.now(timezone.utc)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


async def export_research_memo_pdf(
    content: str,
    *,
    title: str = "Research Memo",
    footnotes: list[dict] | None = None,
) -> bytes:
    """Export a research memo to PDF with footnotes as bibliography."""
    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=1 * inch,
        bottomMargin=1 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
        title=title,
        author="Smriti AI",
        creator="Smriti AI",
        subject="Research Memo",
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        "MemoTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=16,
        alignment=1,
        spaceAfter=20,
    )

    heading_style = ParagraphStyle(
        "MemoHeading",
        parent=styles["Heading1"],
        fontName="Times-Bold",
        fontSize=14,
        spaceBefore=14,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        "MemoBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=12,
        leading=16,
        spaceAfter=6,
    )

    bib_style = ParagraphStyle(
        "MemoBib",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=10,
        leading=13,
        spaceAfter=2,
    )

    flowables: list[Paragraph | Spacer] = []
    flowables.append(Paragraph(title, title_style))
    flowables.append(Spacer(1, 12))

    sections = _parse_sections(content)
    for heading, body_lines in sections:
        if heading is not None:
            flowables.append(Paragraph(heading, heading_style))
        body_text = "\n".join(body_lines).strip()
        if not body_text:
            continue
        for para_text in body_text.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            safe_text = (
                para_text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )
            flowables.append(Paragraph(safe_text, body_style))

    # Bibliography
    if footnotes:
        flowables.append(Spacer(1, 12))
        flowables.append(Paragraph("Bibliography", heading_style))
        for fn in footnotes:
            num = fn.get("number", 0)
            citation = fn.get("citation", "")
            title_text = fn.get("title", "")
            label = f"[{num}] {citation}"
            if title_text and title_text != citation:
                label += f" — {title_text}"
            safe_label = (
                label.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            flowables.append(Paragraph(safe_label, bib_style))

    doc.build(flowables)
    buf.seek(0)
    return buf.read()


# ---------------------------------------------------------------------------
# Filing package export (ZIP)
# ---------------------------------------------------------------------------


async def export_filing_package(
    content: str,
    template: DocumentTemplate,
    *,
    title: str = "",
    court_profile: "CourtProfile | None" = None,
    affidavit: str = "",
    vakalatnama: bool = True,
    annexure_index: list[dict] | None = None,
) -> bytes:
    """Export a complete filing package as a ZIP file.

    Contains: main document (DOCX), affidavit (DOCX if provided),
    vakalatnama template (DOCX), and annexure index (DOCX).
    """
    import zipfile

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # 1. Main document
        docx_bytes = await export_to_docx(
            content, template, title=title,
            court_profile=court_profile,
        )
        safe_name = template.doc_type.replace(" ", "_")
        zf.writestr(f"01_main_{safe_name}.docx", docx_bytes)

        # 2. Affidavit (if provided)
        if affidavit:
            aff_bytes = await export_to_docx(
                affidavit, template, title="Affidavit",
                court_profile=court_profile,
            )
            zf.writestr("02_affidavit.docx", aff_bytes)

        # 3. Vakalatnama template
        if vakalatnama:
            vak_content = _generate_vakalatnama_template(template)
            vak_doc = Document()
            for section in vak_doc.sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            style = vak_doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            title_para = vak_doc.add_heading("VAKALATNAMA", level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in title_para.runs:
                run.font.name = "Times New Roman"
            para = vak_doc.add_paragraph(vak_content)
            para.style.font.name = "Times New Roman"
            vak_buf = io.BytesIO()
            vak_doc.save(vak_buf)
            vak_buf.seek(0)
            zf.writestr("03_vakalatnama.docx", vak_buf.read())

        # 4. Annexure index
        if annexure_index:
            idx_doc = Document()
            style = idx_doc.styles["Normal"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(12)
            idx_doc.add_heading("INDEX OF ANNEXURES", level=0)
            table = idx_doc.add_table(rows=1, cols=3)
            table.style = "Table Grid"
            hdr = table.rows[0].cells
            hdr[0].text = "Sr. No."
            hdr[1].text = "Description"
            hdr[2].text = "Page No."
            for i, ann in enumerate(annexure_index, 1):
                row = table.add_row().cells
                row[0].text = str(i)
                row[1].text = ann.get("description", "")
                row[2].text = ann.get("page", "")
            idx_buf = io.BytesIO()
            idx_doc.save(idx_buf)
            idx_buf.seek(0)
            zf.writestr("04_annexure_index.docx", idx_buf.read())

    buf.seek(0)
    return buf.read()


def _generate_vakalatnama_template(template: DocumentTemplate) -> str:
    """Generate a standard vakalatnama text template."""
    return (
        "I/We, the undersigned, do hereby appoint and retain "
        "Shri/Smt. _________________, Advocate, "
        "as my/our Advocate to appear, act and plead on my/our behalf "
        f"in connection with the {template.display_name} "
        "and to conduct and prosecute (or defend) the same and all "
        "proceedings that may be taken in respect of any application "
        "connected with the same or any decree or order passed therein.\n\n"
        "AND I/We do hereby authorize the said Advocate to sign, verify "
        "and present pleadings, applications, and other documents.\n\n"
        "AND I/We agree to ratify all acts done by the said Advocate "
        "in pursuance of this authority.\n\n"
        "Dated this ___ day of _____________, 20___\n\n"
        "________________________\n"
        "Signature of the Client\n\n"
        "Accepted:\n\n"
        "________________________\n"
        "Advocate\n"
        "Enrollment No.: _______________\n"
    )
